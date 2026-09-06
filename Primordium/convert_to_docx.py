"""Convert Primordium README.md to Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

SRC = "DesignTool/Primordium/README.md"
DST = "DesignTool/Primordium/Primordium_元初.docx"

doc = Document()

# -- style tweaks --
style = doc.styles["Normal"]
style.font.name = "SimSun"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

for level in range(1, 5):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "SimHei"
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "4",
        qn("w:color"): "AAAAAA",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_blockquote(doc, text):
    """Add indented, italic paragraph for blockquotes."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_code_block(doc, lines):
    """Add monospace code block."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def parse_inline(doc, paragraph, text):
    """Parse inline formatting: **bold**, `code`."""
    # Split by ** pairs
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # handle inline code within non-bold segments
            sub_parts = re.split(r'(`[^`]+`)', part)
            for sp in sub_parts:
                if sp.startswith("`") and sp.endswith("`"):
                    run = paragraph.add_run(sp[1:-1])
                    run.font.name = "Courier New"
                    run.font.size = Pt(9.5)
                else:
                    paragraph.add_run(sp)

def add_markdown_table(doc, rows):
    """Add a table from markdown rows (list of lists of cells)."""
    if not rows:
        return
    # rows[0] is header, rows[1] is separator, rest are data
    data_rows = [rows[0]] + rows[2:] if len(rows) > 2 else rows
    ncols = len(data_rows[0])
    table = doc.add_table(rows=len(data_rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(data_rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if i == 0:
                run = p.add_run(cell_text.strip())
                run.bold = True
                run.font.size = Pt(10)
                # header background
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn("w:shd"), {
                    qn("w:fill"): "E8E8E8",
                    qn("w:val"): "clear",
                })
                shading.append(shd)
            else:
                run = p.add_run(cell_text.strip())
                run.font.size = Pt(10)

def add_paragraph(doc, text):
    """Add a normal paragraph with inline formatting."""
    if not text.strip():
        return
    # check if it's a list item
    is_list = re.match(r'^(\d+[.、]|[-*])\s+', text)
    p = doc.add_paragraph()
    if is_list:
        p.paragraph_format.left_indent = Cm(1.0)
        p.style = doc.styles["List Bullet"] if re.match(r'^[-*]', text) else doc.styles["List Number"]
    parse_inline(doc, p, text)
    return p


# -- read and parse --
with open(SRC, "r", encoding="utf-8") as f:
    lines = f.readlines()

in_code = False
code_buf = []
in_table = False
table_buf = []

i = 0
while i < len(lines):
    line = lines[i].rstrip()

    # code block
    if line.startswith("```"):
        if in_code:
            add_code_block(doc, code_buf)
            code_buf = []
            in_code = False
        else:
            in_code = True
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # table
    if "|" in line and line.strip().startswith("|"):
        if not in_table:
            in_table = True
            table_buf = []
        cells = [c.strip() for c in line.strip().split("|")[1:-1]]
        table_buf.append(cells)
        # peek next line - if not a table row, flush
        if i + 1 < len(lines) and "|" not in lines[i + 1]:
            add_markdown_table(doc, table_buf)
            table_buf = []
            in_table = False
        elif i + 1 >= len(lines):
            add_markdown_table(doc, table_buf)
            table_buf = []
            in_table = False
        i += 1
        continue
    else:
        if in_table:
            add_markdown_table(doc, table_buf)
            table_buf = []
            in_table = False

    # blank line
    if not line.strip():
        i += 1
        continue

    # horizontal rule
    if line.strip() == "---":
        add_horizontal_rule(doc)
        i += 1
        continue

    # heading
    h_match = re.match(r'^(#{1,4})\s+(.+)$', line)
    if h_match:
        level = len(h_match.group(1))
        heading_text = h_match.group(2).strip()
        doc.add_heading(heading_text, level=level)
        i += 1
        continue

    # blockquote
    if line.startswith(">"):
        text = line[1:].strip()
        # handle bold in blockquote
        add_blockquote(doc, text)
        i += 1
        continue

    # normal paragraph
    add_paragraph(doc, line)
    i += 1

# flush any trailing table
if table_buf:
    add_markdown_table(doc, table_buf)

# -- page setup --
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

doc.save(DST)
print(f"Done → {DST}")
